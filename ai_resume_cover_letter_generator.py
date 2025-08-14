# ai_resume_cover_letter_generator_scroll.py

import os
import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- Document helper functions ----------
def add_heading_center(doc, text, size=18, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

def add_subheading(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def add_bullets(doc, items):
    for it in items:
        if it.strip():
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(it.strip())

def make_resume(data, out_path):
    doc = Document()

    add_heading_center(doc, data["name"], size=20)

    contact_line = " | ".join(
        [v for v in [
            data.get("email", "").strip(),
            data.get("phone", "").strip(),
            data.get("address", "").strip(),
            data.get("linkedin", "").strip(),
            data.get("github", "").strip(),
            data.get("portfolio", "").strip(),
        ] if v]
    )
    if contact_line:
        p = doc.add_paragraph(contact_line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if data.get("summary", "").strip():
        add_subheading(doc, "Professional Summary")
        doc.add_paragraph(data["summary"].strip())

    skills = [s.strip() for s in data.get("skills", "").split(",") if s.strip()]
    if skills:
        add_subheading(doc, "Skills")
        add_bullets(doc, skills)

    if data.get("education", "").strip():
        add_subheading(doc, "Education")
        edu_lines = [ln.strip() for ln in data["education"].split("\n") if ln.strip()]
        add_bullets(doc, edu_lines)

    if data.get("experience", "").strip():
        add_subheading(doc, "Experience")
        exp_lines = [ln.strip() for ln in data["experience"].split("\n") if ln.strip()]
        add_bullets(doc, exp_lines)

    if data.get("projects", "").strip():
        add_subheading(doc, "Projects")
        proj_lines = [ln.strip() for ln in data["projects"].split("\n") if ln.strip()]
        add_bullets(doc, proj_lines)

    if data.get("certifications", "").strip():
        add_subheading(doc, "Certifications")
        cert_lines = [ln.strip() for ln in data["certifications"].split("\n") if ln.strip()]
        add_bullets(doc, cert_lines)

    doc.save(out_path)

def make_cover_letter(data, out_path):
    today = datetime.date.today().strftime("%B %d, %Y")
    doc = Document()

    doc.add_paragraph(f"{today}\n")
    doc.add_paragraph(f"{data.get('hiring_manager', 'Hiring Manager')}")
    doc.add_paragraph(f"{data.get('company', 'Company Name')}\n")

    doc.add_paragraph(f"Dear {data.get('hiring_manager', 'Hiring Manager')},\n")

    doc.add_paragraph(
        f"I am excited to apply for the {data.get('target_role', '[Role]')} position at {data.get('company', 'your company')}. "
        f"With my skills in {data.get('skills', '')} and experience in software development, "
        "I am confident in my ability to contribute effectively."
    )

    doc.add_paragraph(
        "I have worked on several projects, including:\n" +
        "\n".join([f"- {p}" for p in data.get("projects", "").split("\n") if p.strip()])
    )

    doc.add_paragraph(
        "I look forward to the opportunity to discuss how I can contribute to your team.\n"
        "Thank you for considering my application."
    )

    doc.add_paragraph("\nSincerely,")
    doc.add_paragraph(data.get("name", ""))
    doc.add_paragraph(data.get("email", ""))
    doc.add_paragraph(data.get("phone", ""))

    doc.save(out_path)

# ---------- Main GUI ----------
class ResumeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Resume & Cover Letter Generator - Day 25")
        self.root.state("zoomed")  # open maximized

        # Scrollable frame setup
        canvas = tk.Canvas(root)
        scrollbar = tk.Scrollbar(root, orient="vertical", command=canvas.yview)
        scroll_frame = tk.Frame(canvas)

        scroll_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Fields
        self.entries = {}
        fields = [
            ("Full Name*", "name"),
            ("Email", "email"),
            ("Phone", "phone"),
            ("Address", "address"),
            ("LinkedIn URL", "linkedin"),
            ("GitHub URL", "github"),
            ("Portfolio URL", "portfolio"),
            ("Summary", "summary", True),
            ("Skills (comma-separated)", "skills"),
            ("Education (one per line)", "education", True),
            ("Experience (one per line)", "experience", True),
            ("Projects (one per line)", "projects", True),
            ("Certifications (one per line)", "certifications", True),
            ("Target Role", "target_role"),
            ("Company", "company"),
            ("Hiring Manager", "hiring_manager"),
        ]

        for idx, field in enumerate(fields):
            label_text = field[0]
            key = field[1]
            multiline = len(field) > 2 and field[2]

            tk.Label(scroll_frame, text=label_text).grid(row=idx, column=0, sticky="ne", padx=5, pady=5)
            if multiline:
                text_widget = tk.Text(scroll_frame, width=70, height=5)
                text_widget.grid(row=idx, column=1, padx=5, pady=5)
                self.entries[key] = text_widget
            else:
                entry_widget = tk.Entry(scroll_frame, width=70)
                entry_widget.grid(row=idx, column=1, padx=5, pady=5)
                self.entries[key] = entry_widget

        # Output folder
        tk.Label(scroll_frame, text="Output Folder").grid(row=len(fields), column=0, sticky="e", padx=5, pady=5)
        self.output_dir = tk.StringVar(value=os.path.join(os.getcwd(), "Day25_Output"))
        tk.Entry(scroll_frame, textvariable=self.output_dir, width=50).grid(row=len(fields), column=1, sticky="w", padx=5, pady=5)
        tk.Button(scroll_frame, text="Browse", command=self.browse_dir).grid(row=len(fields), column=2, padx=5, pady=5)

        # Submit button
        tk.Button(
            scroll_frame,
            text="Generate Resume & Cover Letter",
            command=self.generate_files,
            bg="#4CAF50", fg="white", font=("Arial", 12, "bold")
        ).grid(row=len(fields) + 1, column=1, pady=20)

    def browse_dir(self):
        path = filedialog.askdirectory()
        if path:
            self.output_dir.set(path)

    def collect_data(self):
        data = {}
        for key, widget in self.entries.items():
            if isinstance(widget, tk.Text):
                data[key] = widget.get("1.0", "end").strip()
            else:
                data[key] = widget.get().strip()
        return data

    def generate_files(self):
        data = self.collect_data()

        if not data["name"]:
            messagebox.showerror("Error", "Full Name is required.")
            return

        out_dir = self.output_dir.get()
        os.makedirs(out_dir, exist_ok=True)

        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        run_dir = os.path.join(out_dir, f"Resume_Cover_{timestamp}")
        os.makedirs(run_dir, exist_ok=True)

        resume_path = os.path.join(run_dir, "Resume.docx")
        cover_path = os.path.join(run_dir, "Cover_Letter.docx")

        make_resume(data, resume_path)
        make_cover_letter(data, cover_path)

        messagebox.showinfo("Success", f"Files generated:\n{resume_path}\n{cover_path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = ResumeApp(root)
    root.mainloop()
